"""
Generates the Norvane Operator Playbook as a styled Word document.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colour palette ────────────────────────────────────────────────────────────
CHARCOAL   = RGBColor(0x1A, 0x1A, 0x2E)   # near-black navy
SLATE      = RGBColor(0x16, 0x21, 0x3E)   # section header bg
ACCENT     = RGBColor(0x0F, 0x7B, 0x6C)   # teal accent
ACCENT_MID = RGBColor(0x1A, 0x99, 0x88)   # lighter teal for table headers
LIGHT_BG   = RGBColor(0xF4, 0xF7, 0xF6)   # pale mint table rows
MID_GREY   = RGBColor(0x6B, 0x7B, 0x8D)   # body text grey
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
WARN_AMBER = RGBColor(0xE0, 0x7B, 0x39)   # "avoid / don't" callouts
DIVIDER    = RGBColor(0xD0, 0xD8, 0xD5)

# ── Helpers ───────────────────────────────────────────────────────────────────

def rgb_hex(colour: RGBColor) -> str:
    # RGBColor is a 3-tuple subclass; __str__ returns the 6-char hex
    return str(colour).upper()


def set_cell_bg(cell, colour: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  rgb_hex(colour))
    tcPr.append(shd)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, colour in [("top", top), ("bottom", bottom),
                          ("left", left), ("right", right)]:
        if colour:
            el = OxmlElement(f"w:{side}")
            hex_col = rgb_hex(colour)
            el.set(qn("w:val"),   "single")
            el.set(qn("w:sz"),    "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), hex_col)
            tcBorders.append(el)
    tcPr.append(tcBorders)


def add_paragraph_border_bottom(para, colour: RGBColor):
    """Add a bottom border to a paragraph (used for section rule)."""
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    hex_col = rgb_hex(colour)
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), hex_col)
    pBdr.append(bot)
    pPr.append(pBdr)


def set_para_shading(para, colour: RGBColor):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    hex_col = rgb_hex(colour)
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_col)
    pPr.append(shd)


def page_break(doc):
    doc.add_page_break()


def set_col_width(table, col_idx, width_inches):
    for row in table.rows:
        row.cells[col_idx].width = Inches(width_inches)


# ── Base style shortcuts ──────────────────────────────────────────────────────

def body_run(para, text, bold=False, italic=False, colour=None, size=Pt(10.5)):
    run = para.add_run(text)
    run.font.size  = size
    run.font.bold  = bold
    run.font.italic = italic
    run.font.color.rgb = colour or MID_GREY
    run.font.name  = "Calibri"
    return run


def add_body(doc, text, bold=False, italic=False, colour=None,
             space_before=Pt(2), space_after=Pt(5), left_indent=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = space_before
    para.paragraph_format.space_after  = space_after
    if left_indent:
        para.paragraph_format.left_indent = left_indent
    body_run(para, text, bold=bold, italic=italic, colour=colour)
    return para


def add_bullet(doc, text, level=0, colour=None):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(3)
    para.paragraph_format.left_indent  = Inches(0.25 + level * 0.25)
    body_run(para, text, colour=colour or MID_GREY)
    return para


# ── Section heading ───────────────────────────────────────────────────────────

def add_section_heading(doc, number: str, title: str):
    """Dark navy pill-style section heading."""
    # Spacer
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(16)
    sp.paragraph_format.space_after  = Pt(0)

    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(8)
    set_para_shading(para, SLATE)

    # Section number chip
    r1 = para.add_run(f"  {number}  ")
    r1.font.name  = "Calibri"
    r1.font.size  = Pt(9)
    r1.font.bold  = True
    r1.font.color.rgb = ACCENT_MID

    # Title
    r2 = para.add_run(title.upper() + "  ")
    r2.font.name  = "Calibri"
    r2.font.size  = Pt(11)
    r2.font.bold  = True
    r2.font.color.rgb = WHITE

    return para


def add_h2(doc, text):
    """Teal underlined sub-heading."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after  = Pt(4)
    r = para.add_run(text)
    r.font.name  = "Calibri"
    r.font.size  = Pt(11)
    r.font.bold  = True
    r.font.color.rgb = ACCENT
    add_paragraph_border_bottom(para, ACCENT)
    return para


def add_h3(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(7)
    para.paragraph_format.space_after  = Pt(3)
    r = para.add_run(text)
    r.font.name  = "Calibri"
    r.font.size  = Pt(10.5)
    r.font.bold  = True
    r.font.color.rgb = CHARCOAL
    return para


# ── Callout box ───────────────────────────────────────────────────────────────

def add_callout(doc, label: str, lines: list, style="info"):
    """A single-cell table used as a callout / pull-quote."""
    bg = LIGHT_BG if style == "info" else RGBColor(0xFF, 0xF3, 0xE8)
    border_col = ACCENT if style == "info" else WARN_AMBER

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg)
    set_cell_borders(cell, left=border_col)
    cell.width = Inches(6.3)

    # Label
    lp = cell.paragraphs[0]
    lp.paragraph_format.space_before = Pt(4)
    lp.paragraph_format.space_after  = Pt(2)
    lr = lp.add_run(label)
    lr.font.name  = "Calibri"
    lr.font.size  = Pt(9)
    lr.font.bold  = True
    lr.font.color.rgb = border_col

    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(line)
        r.font.name  = "Calibri"
        r.font.size  = Pt(10)
        r.font.color.rgb = CHARCOAL

    # Bottom spacer
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


# ── Data table ────────────────────────────────────────────────────────────────

def add_data_table(doc, headers: list, rows: list, col_widths=None):
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"

    # Header row
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, ACCENT)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(h)
        r.font.name  = "Calibri"
        r.font.size  = Pt(9.5)
        r.font.bold  = True
        r.font.color.rgb = WHITE

    # Data rows
    for ri, row_data in enumerate(rows):
        tr = tbl.rows[ri + 1]
        bg = LIGHT_BG if ri % 2 == 0 else WHITE
        for ci, cell_text in enumerate(row_data):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            # Bold first cell of each row
            bold = (ci == 0)
            r = p.add_run(cell_text)
            r.font.name  = "Calibri"
            r.font.size  = Pt(9.5)
            r.font.bold  = bold
            r.font.color.rgb = CHARCOAL

    if col_widths:
        for ci, w in enumerate(col_widths):
            set_col_width(tbl, ci, w)

    # Spacer after table
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)
    return tbl


# ── Flow / chain block ────────────────────────────────────────────────────────

def add_flow(doc, steps: list):
    """Vertical step flow with teal numbers."""
    for i, step in enumerate(steps):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after  = Pt(2)
        para.paragraph_format.left_indent  = Inches(0.2)

        num = para.add_run(f"  {i+1:02d}  ")
        num.font.name  = "Calibri"
        num.font.size  = Pt(9)
        num.font.bold  = True
        num.font.color.rgb = WHITE
        # Fake chip: we can't do inline shading in Word paragraphs easily,
        # so we'll just use bold teal for the number
        num.font.color.rgb = ACCENT

        txt = para.add_run(f"  {step}")
        txt.font.name  = "Calibri"
        txt.font.size  = Pt(10.5)
        txt.font.color.rgb = CHARCOAL

        if i < len(steps) - 1:
            arr = doc.add_paragraph()
            arr.paragraph_format.left_indent  = Inches(0.3)
            arr.paragraph_format.space_before = Pt(0)
            arr.paragraph_format.space_after  = Pt(0)
            ar = arr.add_run("↓")
            ar.font.size  = Pt(10)
            ar.font.color.rgb = ACCENT_MID


# ═══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT BUILD
# ═══════════════════════════════════════════════════════════════════════════════

def build():
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ─────────────────────────────────────────────────────────────────────────
    # COVER PAGE
    # ─────────────────────────────────────────────────────────────────────────
    doc.add_paragraph()  # top space

    # Company name
    co = doc.add_paragraph()
    co.alignment = WD_ALIGN_PARAGRAPH.CENTER
    co.paragraph_format.space_after = Pt(4)
    cr = co.add_run("NORVANE")
    cr.font.name  = "Calibri"
    cr.font.size  = Pt(36)
    cr.font.bold  = True
    cr.font.color.rgb = ACCENT

    # Tagline
    tl = doc.add_paragraph()
    tl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tl.paragraph_format.space_after = Pt(24)
    tr_ = tl.add_run("OPERATOR PLAYBOOK")
    tr_.font.name  = "Calibri"
    tr_.font.size  = Pt(16)
    tr_.font.bold  = False
    tr_.font.color.rgb = CHARCOAL

    # Horizontal rule (via paragraph border)
    rule = doc.add_paragraph("  ")
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph_border_bottom(rule, ACCENT)
    rule.paragraph_format.space_after = Pt(24)

    # Sub-title block
    for line, sz, col in [
        ("Solo Founder GTM Field Manual", 13, MID_GREY),
        ("Go-to-Market · Validation · Sales · Operations", 11, MID_GREY),
        ("May 2026", 10, RGBColor(0xAA, 0xBB, 0xB8)),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(line)
        r.font.name  = "Calibri"
        r.font.size  = Pt(sz)
        r.font.color.rgb = col

    doc.add_paragraph()

    add_callout(doc,
        "PURPOSE OF THIS DOCUMENT",
        [
            "This playbook is an operational execution system — not a strategy deck.",
            "It covers: 7-day GTM execution, founder-led sales, validation frameworks,",
            "outreach scripts, objection handling, and what NOT to spend time on.",
            "Every section is designed to produce a learning or a revenue event within 14 days.",
        ],
        style="info"
    )

    page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # FRAME
    # ─────────────────────────────────────────────────────────────────────────
    add_section_heading(doc, "00", "FRAME — WHAT YOU'RE ACTUALLY VALIDATING")

    add_body(doc,
        "Before any tactics, be precise about what validation means for Norvane. "
        "Confusing intellectual validation with market validation is the most common early-stage mistake.",
        space_after=Pt(8)
    )

    add_h2(doc, "You Are NOT Validating")
    for line in [
        "Whether operational fragmentation is a real problem (it obviously is)",
        "Whether coordination tools are valuable in theory",
        "Whether SMEs in ag/logistics have pain (they do)",
    ]:
        add_bullet(doc, line)

    add_h2(doc, "You ARE Validating")
    for line in [
        "Whether a specific person will pay for a diagnostic engagement",
        "Whether your framing creates recognition before you explain the solution",
        "Whether you can reach decision-makers efficiently via direct outreach",
        "Whether the diagnostic format matches how these buyers make decisions",
    ]:
        add_bullet(doc, line)

    add_callout(doc,
        "CRITICAL DISTINCTION",
        ['"Yes, this sounds interesting" is not a signal.',
         '"How much does this cost?" is a signal.',
         'Design your outreach and calls to reach the second statement, not the first.'],
        style="warn"
    )

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 1 — PRIORITISATION
    # ─────────────────────────────────────────────────────────────────────────
    add_section_heading(doc, "01", "RUTHLESS PRIORITISATION")

    add_h2(doc, "The Dependency Chain for Your First Dollar")

    add_flow(doc, [
        "Identify the right person with active operational pain",
        "Reach them directly (LinkedIn DM or email)",
        "Get a 20-minute discovery call",
        "Listen to their operational reality",
        "Propose a diagnostic engagement (fixed scope, fixed price)",
        "Get a yes or a no — both are data",
    ])

    add_body(doc, "\nEvery hour not spent on that chain is waste. The website is done enough. "
             "The positioning is close enough. You need conversations, not polish.", space_before=Pt(8))

    add_h2(doc, "Three Things That Will Kill Momentum in Week 1")
    for line in [
        "Website improvements instead of outreach — this is productive procrastination",
        "Building a deck or methodology document before you've had 5 real conversations",
        "Waiting for LinkedIn to warm up before sending cold messages",
    ]:
        add_bullet(doc, line, colour=WARN_AMBER)

    add_h2(doc, "Correct Order of Operations")
    add_flow(doc, [
        "Outreach",
        "Calls",
        "Diagnose what they actually need",
        "Build the minimum thing that delivers it",
        "Charge for it",
    ])

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 2 — 7-DAY PLAN
    # ─────────────────────────────────────────────────────────────────────────
    add_section_heading(doc, "02", "7-DAY EXECUTION PLAN")

    days = [
        ("Day 1", "Target List Build (4 hours)",
         ["Build 50-person LinkedIn outreach list (logistics-first: ops managers, owners, GMs)",
          "Target: 10-200 employees, operationally hands-on, seasonal or multi-site complexity",
          "Write three outreach message variants (see Section 08)",
          "Send first 15 messages — 5 per variant to A/B test"]),
        ("Day 2", "First Outreach Wave + LinkedIn",
         ["Follow up on Day 1 non-replies with a single line",
          "Publish first LinkedIn post (operational scenario + question — no mention of Norvane)",
          "Set up Calendly free tier or Google Meet link"]),
        ("Day 3", "Follow-up + Call Prep",
         ["Send second outreach wave (15 messages, same A/B structure)",
          "Write call script and read it aloud once — time it (target: 20 minutes)",
          "Create 5-question intake form in Google Forms"]),
        ("Day 4", "Second Wave + First Calls",
         ["Send 10-15 more DMs, refined based on which variant got replies",
          "Take every call that has booked — do not reschedule",
          "If no calls yet: start email outreach on the same list"]),
        ("Day 5", "Pattern Synthesis",
         ["Stop outreach. Review every response — replies, objections, silence",
          "Which job title replied most? Which industry? What pain language did they use?",
          "Did anyone ask 'how much?' — strong buying signal",
          "Update targeting and messaging before next wave"]),
        ("Day 6", "Third Wave + Second Post",
         ["Send 15-20 refined messages based on Day 5 synthesis",
          "Try email if LinkedIn is underperforming",
          "Publish second LinkedIn post: operational scenario with cost framing + question"]),
        ("Day 7", "Weekly Review + Week 2 Plan",
         ["Conversations started: target 5+",
          "Calls booked: target 2+ | Calls completed: target 1+",
          "Any 'how much does this cost' questions: yes/no",
          "If 0 conversations after 50 outreaches: message or targeting is wrong — diagnose before continuing"]),
    ]

    for day, subtitle, bullets in days:
        add_h2(doc, f"{day} — {subtitle}")
        for b in bullets:
            add_bullet(doc, b)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 3 — DAILY TARGETS
    # ─────────────────────────────────────────────────────────────────────────
    add_section_heading(doc, "03", "DAILY OUTREACH TARGETS")

    add_data_table(doc,
        ["Day", "New Outreach", "Follow-ups", "Call Target", "Note"],
        [
            ["1", "15", "0",  "0",   "3 variants × 5"],
            ["2", "0",  "5",  "0",   "Let wave 1 breathe"],
            ["3", "15", "10", "0–1", "Refine from wave 1 replies"],
            ["4", "0",  "10", "0–1", "Add email if LinkedIn flat"],
            ["5", "15", "10", "1–2", "Synthesise before sending"],
            ["6", "10", "10", "1–2", "Refined targeting"],
            ["7", "0",  "5",  "1–2", "Review + Week 2 plan"],
        ],
        col_widths=[0.5, 1.1, 1.1, 1.1, 3.0]
    )

    add_body(doc,
        "Total: ~55 new outreaches, ~50 follow-ups, target 3–5 calls booked in Week 1. "
        "Enough volume to generate signal, not so much that personalization disappears. "
        "In this segment, mass outreach without personalization tanks response rates."
    )

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 4 — VALIDATION STRATEGY
    # ─────────────────────────────────────────────────────────────────────────
    add_section_heading(doc, "04", "FASTEST VALIDATION STRATEGY")

    add_body(doc,
        "The only valid validation for Norvane in the next 30 days is a paid diagnostic engagement. "
        "Not LinkedIn followers. Not website visits. Not 'people who said they'd be interested next quarter.'",
        bold=False, space_after=Pt(8)
    )

    add_h2(doc, "Validation Signal Ladder")

    add_data_table(doc,
        ["Signal", "Strength", "What It Actually Means"],
        [
            ["Someone pays for diagnostic",              "Very High", "Real demand"],
            ["Someone asks 'how much' or 'when can we start'", "High", "Buying intent"],
            ["Prospect describes pain in your language unprompted", "High", "Positioning resonating"],
            ["Prospect refers you to a colleague",       "Medium-High", "Offer is understandable"],
            ["'We have this exact problem'",             "Medium",    "Pain confirmed, solution not yet"],
            ["'This sounds interesting'",                "Low",       "Polite noise"],
            ["LinkedIn like / website visit",            "Very Low",  "Vanity — ignore"],
        ],
        col_widths=[2.8, 1.3, 2.6]
    )

    add_h2(doc, "The 7–14 Day Validation Method")
    for line in [
        "Run direct outreach at the cadence in Section 03",
        "Take every discovery call that comes from outreach",
        "In each call: describe the diagnostic engagement, name a price point, track the reaction",
        "If they don't ask 'how does that work' or 'how much' after your offer — positioning is not landing",
    ]:
        add_bullet(doc, line)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 5 — MVP FEATURE
    # ─────────────────────────────────────────────────────────────────────────
    add_section_heading(doc, "05", "MINIMUM VIABLE OPERATIONAL FEATURE")

    add_callout(doc,
        "IMPORTANT",
        ["Do not build software right now.",
         "Norvane Phase 1 is a consulting engagement. The minimum viable feature is the diagnostic process, delivered manually.",
         "No software required. No platform. No dashboard."],
        style="warn"
    )

    add_h2(doc, "The Manual Diagnostic Process")

    steps_detail = [
        ("Pre-call Intake Form (Google Form, 5 questions)",
         ["How many people are in your operations team?",
          "What is the biggest source of information loss in your operation?",
          "How do you currently track job status or task completion?",
          "What decision do you most often make without the right information?",
          "What would it mean to your operation if you could fix that?"]),
        ("Discovery Call (60 min, Loom-recorded with permission)",
         ["Use the call script in Section 15",
          "Take notes in a structured template immediately after"]),
        ("Operations Map (Google Slides or Figma, 1 page)",
         ["Current information flows",
          "Decision bottlenecks",
          "Where handoffs break down",
          "What is invisible to leadership"]),
        ("Diagnostic Report (PDF, 4–6 pages)",
         ["Current state: what is fragmented and why",
          "Top 3 operational blind spots",
          "Estimated cost of current coordination failures",
          "Recommendations: what to fix first and how"]),
    ]

    for title, bullets in steps_detail:
        add_h3(doc, title)
        for b in bullets:
            add_bullet(doc, b, level=1)

    add_body(doc,
        "This IS the product. It takes 4–6 hours to produce. Charge $1,500–$3,500 for it.",
        bold=True, colour=ACCENT
    )

    add_h2(doc, "When to Build Software")
    add_body(doc,
        "After 3–5 paid diagnostics, you will see patterns. If the same tool keeps being "
        "the correct recommendation, that is when you consider building it. Building before "
        "that point is a bet, not a product."
    )

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 6 — MESSAGING
    # ─────────────────────────────────────────────────────────────────────────
    add_section_heading(doc, "06", "MESSAGING — RECOGNITION OVER HYPE")

    add_body(doc,
        "Most consultants describe their capability. The buyer does not care about your capability. "
        "They care about whether you understand their specific problem better than they currently do."
    )

    add_h2(doc, "The Recognition Test")
    add_body(doc,
        "Your prospect should be able to read your message and think 'yes, that is exactly "
        "what happens here' before you have mentioned your solution. That is the bar."
    )

    add_h2(doc, "Core Message Variants")

    messages = [
        ("Variant A — Pain-First (Operations Language)",
         "Your operations team has the information they need to make good decisions — "
         "it just lives in three different places, none of which talk to each other. "
         "By the time you have a clear picture, the moment to act has already passed."),
        ("Variant B — Cost-Focused",
         "Every day your operation runs on fragmented information, you are making slower "
         "decisions with higher coordination overhead. That does not show up on a P&L line — "
         "but it shows up in margin, in driver turnover, and in harvest windows you cannot recover."),
        ("Variant C — Specificity Play",
         "Between the haul sheet, the field app, and the WhatsApp thread — your operation "
         "is being managed across four systems that were never designed to work together. "
         "The cost is not the software. The cost is what gets missed in the gaps."),
    ]

    for label, text in messages:
        add_h3(doc, label)
        add_callout(doc, "MESSAGE", [text], style="info")

    add_h2(doc, "What to Avoid in All Messaging")
    avoid = [
        "'AI-powered' — sounds like a vendor, not a consultant",
        "'Operational excellence' — too corporate, means nothing to a 15-person ag company",
        "'Coordination intelligence' — too abstract for your actual buyers",
        "'Real-time visibility' — same pitch as every dashboard they have already seen",
        "Fabricated case studies, statistics, or client names",
    ]
    for a in avoid:
        add_bullet(doc, a, colour=WARN_AMBER)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 7 — LINKEDIN
    # ─────────────────────────────────────────────────────────────────────────
    add_section_heading(doc, "07", "LINKEDIN GROWTH STRATEGY")

    add_body(doc,
        "LinkedIn is a credibility channel, not your primary sales channel for Weeks 1–4. "
        "Its job right now: make your DMs more credible. Post twice a week. Do not optimize "
        "for follower counts before you have paying clients."
    )

    add_h2(doc, "Post Type 1 — The Operational Scenario")
    add_callout(doc, "SAMPLE POST", [
        "A grain elevator manager once told me: 'By the time I know we have a scheduling problem,",
        "I have already promised three deliveries I cannot keep.'",
        "",
        "That is not a software problem. It is a visibility problem that looks like a scheduling problem.",
        "",
        "How does your team know when a commitment is at risk — before it is already at risk?",
    ], style="info")

    add_h2(doc, "Post Type 2 — The Counterintuitive Observation")
    add_callout(doc, "SAMPLE POST", [
        "Most operations teams think they need better software.",
        "",
        "What they usually need first is a clear picture of what information they are losing",
        "between systems — and where the decision gaps are.",
        "",
        "Software is often the wrong answer to the right problem.",
    ], style="info")

    add_h2(doc, "Post Type 3 — The Diagnostic Question")
    add_callout(doc, "SAMPLE POST", [
        "Quick operational health check:",
        "",
        "When a field crew finishes a job, how long until that information reaches the person",
        "who plans tomorrow's schedule?",
        "",
        "If the answer is 'a few hours' or 'tomorrow morning' —",
        "that is a coordination gap, not a workflow problem.",
    ], style="info")

    add_h2(doc, "What NOT to Post")
    for line in [
        "Case studies you do not have yet",
        "Statistics you fabricated",
        "Anything that sounds like ad copy for a software product",
        "'Excited to announce' anything until there is something real to announce",
    ]:
        add_bullet(doc, line, colour=WARN_AMBER)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 8 — DIRECT OUTREACH
    # ─────────────────────────────────────────────────────────────────────────
    add_section_heading(doc, "08", "DIRECT OUTREACH STRATEGY")

    add_h2(doc, "Platform Priority")
    add_data_table(doc,
        ["Priority", "Channel", "Why"],
        [
            ["1", "LinkedIn DM",  "Highest response rate for owner-operators and ops leads"],
            ["2", "Email",        "If found via LinkedIn profile or company website"],
            ["3", "Cold call",    "Only for small ag operations where owner answers the phone"],
        ],
        col_widths=[0.7, 1.5, 4.5]
    )

    templates = [
        ("Template A — Pain Recognition (ops managers)",
         ["Hi [Name], I work with operations teams in [ag/logistics/field services]",
          "who are managing coordination across multiple systems — and losing visibility at the handoffs.",
          "",
          "Is that a pattern you deal with? I would like to understand how it shows up",
          "in your operation specifically."]),
        ("Template B — Specificity Play (complex operational profiles)",
         ["Hi [Name] — saw you are managing [X trucks / seasonal harvest / field crews] at [Company].",
          "",
          "I consult specifically on operational visibility gaps — the places where information exists",
          "but does not reach the people making decisions in time.",
          "",
          "Worth a quick call to see if that is relevant to where you are?"]),
        ("Template C — Post Engagement (after they engaged with your content)",
         ["Thanks for the comment on [post] — sounds like you are dealing with exactly this.",
          "",
          "I run short diagnostic engagements to map where an operation's information gaps are.",
          "Not a sales pitch — 20 minutes to see if it is relevant. Would you be up for a call?"]),
        ("Cold Email Template",
         ["Subject: Operations visibility — [Company name]",
          "",
          "Hi [Name],",
          "",
          "I consult with ag and logistics operators on a specific problem: the information their",
          "team generates every day does not reach the right people in time to change decisions.",
          "",
          "Not a software pitch — I do short diagnostic engagements to map where the gaps are",
          "and what they are actually costing.",
          "",
          "Worth a 20-minute call to see if it is relevant to [Company]?",
          "",
          "George | Norvane"]),
    ]

    for label, lines in templates:
        add_h3(doc, label)
        add_callout(doc, "TEMPLATE", lines, style="info")

    add_h2(doc, "Follow-up Sequence")
    add_data_table(doc,
        ["Touchpoint", "Timing", "Action"],
        [
            ["First message",    "Day 1",  "Send one of the three variants"],
            ["Follow-up",        "Day 4",  "One line: 'Quick follow-up — worth a conversation?'"],
            ["Final message",    "Day 8",  "'Happy to circle back if timing is off.'"],
            ["Re-contact",       "Day 60+","Mark and revisit — do not chase past three touches"],
        ],
        col_widths=[1.6, 1.0, 4.1]
    )

    add_h2(doc, "What NOT to Do")
    for line in [
        "Send a wall of text explaining what Norvane is",
        "Lead with 'I am a consultant' — generic opener kills response rate",
        "Ask for 45-60 minute calls before establishing any credibility",
        "Follow up more than twice — in this segment, persistence past two contacts damages you",
    ]:
        add_bullet(doc, line, colour=WARN_AMBER)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 9 — DEMO STRUCTURE
    # ─────────────────────────────────────────────────────────────────────────
    add_section_heading(doc, "09", "DEMO STRUCTURE")

    add_callout(doc, "IMPORTANT FRAMING", [
        "At this stage, your 'demo' is a discovery conversation, not a product walkthrough.",
        "You do not have a product to demo. That is correct and appropriate for a consultancy in Week 1.",
        "If a prospect asks to 'see the tool' before you have done a diagnostic, correct the framing directly.",
    ], style="warn")

    add_h3(doc, "When asked 'Can I see the product?'")
    add_callout(doc, "RESPONSE",
        ["We do not have a generic platform to show — the diagnostic produces a custom view of your",
         "operation. The value is in the mapping, not a demo. That is why I want to start with a",
         "conversation about your actual setup."],
        style="info"
    )

    add_h2(doc, "Call-Based Demo Flow")
    demo_steps = [
        ("Minutes 0–5: Diagnostic Overview",
         "Explain what the engagement is, what the output is, and how long it takes."),
        ("Minutes 5–10: Live Illustration",
         "Use a fictional but operationally accurate example from their industry. "
         "Walk through what a diagnostic finding looks like."),
        ("Minutes 10–20: Ask About Their Operation",
         "'Before I go further — how does this map to what you are actually dealing with?' "
         "Then stop talking. Let them describe their reality."),
        ("If fit is confirmed: The Offer",
         "Name the output, the timeline, and the price. Then ask: "
         "'Does that feel relevant to where you are right now?'"),
    ]

    for step, desc in demo_steps:
        add_h3(doc, step)
        add_body(doc, desc, colour=MID_GREY)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 10 — OBJECTION HANDLING
    # ─────────────────────────────────────────────────────────────────────────
    add_section_heading(doc, "10", "OBJECTION HANDLING")

    objections = [
        ("'We already have systems for that.'",
         ["'What are you using? Most operations teams I work with have the tools — the gap is usually",
          "whether the right information reaches the right person at the right time.",
          "Which system handles that for you?'",
          "",
          "(Then let them talk. They will usually explain why their system does not actually work.)"]),
        ("'We do not have budget for consulting right now.'",
         ["'Understood. When I work through this with operations teams, the cost of a diagnostic is",
          "usually less than one week of coordination overhead. Genuinely curious — what is the",
          "highest-friction part of your current operation? If it is something I can diagnose,",
          "it might be worth revisiting.'"]),
        ("'How is this different from a management consultant?'",
         ["'Traditional consultants give you a process framework. We give you an operational map —",
          "specific to how your business actually works, not how best practice says it should.",
          "The diagnostic takes 2–3 weeks, not 6 months, and the output is actionable, not theoretical.'"]),
        ("'I would need to discuss with my partner/team.'",
         ["'Of course. Can I send you a one-pager on what the diagnostic covers so you have",
          "something concrete to share? And would it make sense to get them on a follow-up call?'"]),
        ("'Not the right time.'",
         ["'What would make it the right time? Sometimes that means something specific —",
          "a busy season, a pending system decision, a headcount change — and if I understand that,",
          "I can come back when it is actually useful to you.'"]),
    ]

    for obj, resp in objections:
        add_h3(doc, obj)
        add_callout(doc, "RESPONSE", resp, style="info")

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 11 — PAIN PATTERN LOG
    # ─────────────────────────────────────────────────────────────────────────
    add_section_heading(doc, "11", "COLLECTING OPERATIONAL PAIN PATTERNS")

    add_body(doc,
        "Build a structured log from the first call forward. This is your product research engine. "
        "Record immediately after every call — takes 5 minutes and compounds into your methodology."
    )

    add_h2(doc, "Call Log — Required Fields")
    add_data_table(doc,
        ["Field", "What to Capture"],
        [
            ["Date + Industry",          "ag / logistics / field ops / other"],
            ["Company size",             "Employees, revenue range if disclosed"],
            ["Job title",                "Who you spoke to"],
            ["Exact quote",              "Their pain in their own words — verbatim"],
            ["Current systems",          "What tools they are already using"],
            ["Decision gap",             "What decision they cannot make fast enough"],
            ["Their framing of problem", "What they think the problem is"],
            ["Your diagnosis",           "What the actual problem seems to be — may differ"],
            ["Interest level",           "1–5"],
            ["Asked about price?",       "Yes / No"],
            ["Follow-up action",         "Concrete next step or date"],
        ],
        col_widths=[2.2, 4.5]
    )

    add_h2(doc, "Pattern Signals to Watch For")
    for line in [
        "Same tool named as part of the problem across 3+ calls (e.g., 'WhatsApp is where everything dies') — product insight",
        "Same job title showing the most pain — that is your primary buyer",
        "Same operational scenario across different industries — universal problem worth building around",
        "Same wrong assumption about the problem — that is your positioning clarity",
    ]:
        add_bullet(doc, line)

    add_body(doc,
        "After 10–15 calls, run a 2-hour synthesis session. Group the pain quotes. "
        "Find the 2–3 problems that come up unprompted. Those are your product."
    )

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 12 — REAL vs FAKE DEMAND
    # ─────────────────────────────────────────────────────────────────────────
    add_section_heading(doc, "12", "REAL DEMAND vs. SUPERFICIAL INTEREST")

    add_h2(doc, "Signals of Real Demand")
    add_data_table(doc,
        ["Signal", "What It Means"],
        [
            ["They ask 'how much does this cost'",                      "Buying intent"],
            ["They describe the problem before you explain it",         "Positioning resonating"],
            ["They reference a specific recent incident that cost them", "Pain is active, not theoretical"],
            ["They ask 'when could you start'",                         "Making space for this"],
            ["They introduce you to a colleague",                       "Believe it is relevant"],
            ["Price pushback but they stay engaged",                    "Cost sensitivity, not lack of demand"],
            ["Detailed questions about the output",                     "Evaluating fit, not being polite"],
        ],
        col_widths=[3.5, 3.2]
    )

    add_h2(doc, "Signals of Superficial Interest")
    add_data_table(doc,
        ["Signal", "What It Means"],
        [
            ["'This sounds really interesting' with no follow-up question", "Polite. Not a signal."],
            ["'I will share this with my team' then silence",               "Exit phrase."],
            ["They agree with everything you say",                          "Managing you, not engaging."],
            ["'We might need something like this eventually'",              "Not now = not real."],
            ["High-energy call, no agreed next step",                       "They enjoyed the chat, not the offer."],
            ["Ask for a proposal/deck without discussing specifics",         "Window-shopping."],
        ],
        col_widths=[3.5, 3.2]
    )

    add_h2(doc, "The One-Question Test")
    add_callout(doc, "END EVERY CALL WITH THIS",
        ["'Based on what we have talked about — does a diagnostic engagement feel relevant",
         "to where you are right now, or is there a reason it is not the right fit?'",
         "",
         "Then stop talking. Their response tells you everything.",
         "",
         "'Yes, relevant — can you send a proposal?' = Real interest.",
         "'It is relevant but I need to check budget' = Real interest.",
         "'Yes, interesting' with no action = Superficial interest."],
        style="info"
    )

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 13 — FAKE VALIDATION
    # ─────────────────────────────────────────────────────────────────────────
    add_section_heading(doc, "13", "AVOIDING FAKE VALIDATION")

    add_body(doc,
        "The most dangerous thing at this stage is interpreting encouragement as validation. "
        "Analytically-minded founders are especially vulnerable to this because intellectual "
        "validation feels like market validation but produces completely different outcomes."
    )

    add_h2(doc, "The Five Forms of Fake Validation")
    fake = [
        ("The Nodding Expert",
         "Someone with domain expertise says 'yes, this is a real problem' — but they are not a buyer. "
         "Intellectual validation is not market validation."),
        ("The Interest vs. Purchase Gap",
         "50 people say the website looks great. Zero have paid. "
         "That is 50 data points of noise, not validation."),
        ("The Feature Request Trap",
         "People on calls ask for specific features. This feels like demand. "
         "It is often wishful thinking projected onto your product."),
        ("The Wrong ICP Problem",
         "Enthusiastic responses from people who cannot buy — junior staff, enterprise procurement, wrong industries. "
         "Enthusiasm from the wrong person is misleading."),
        ("The Conference/LinkedIn Trap",
         "Startup community engagement has zero correlation with buyer behavior."),
    ]
    for title, desc in fake:
        add_h3(doc, title)
        add_body(doc, desc, colour=MID_GREY)

    add_h2(doc, "The Only Real Validation Tests")
    for line in [
        "Test 1: Did someone pay for the diagnostic?",
        "Test 2: Did someone agree to pay and only not pay due to a specific, fixable reason (timing, budget cycle, not the decision-maker)?",
        "Test 3: Did someone describe your service accurately to a colleague without being prompted?",
    ]:
        add_bullet(doc, line, colour=ACCENT)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 14 — INDUSTRY PRIORITY
    # ─────────────────────────────────────────────────────────────────────────
    add_section_heading(doc, "14", "HOW TO PRIORITISE INDUSTRIES")

    add_body(doc,
        "Do not spread across three industries equally. Pick one primary for the first 30 days, "
        "then expand based on where you actually get traction."
    )

    add_h2(doc, "Industry Scoring Matrix")
    add_data_table(doc,
        ["Factor", "Agriculture", "Logistics (Trucking)", "Field Operations"],
        [
            ["LinkedIn density",        "Medium",          "High",              "Medium-Low"],
            ["Pain urgency",            "Seasonal (high)", "Always-on",         "Variable"],
            ["Cold outreach accessibility", "High",        "Medium (guarded)",  "Medium"],
            ["Ability to pay $1.5-3.5K", "Medium",        "High",              "Variable"],
            ["Referral density",        "High",            "Medium",            "Low"],
            ["Sales cycle speed",       "Medium",          "Fast",              "Medium"],
        ],
        col_widths=[2.5, 1.5, 1.7, 1.7]
    )

    add_callout(doc, "RECOMMENDATION",
        ["Primary industry for Days 1–30: Logistics (trucking, regional freight)",
         "",
         "Why: year-round urgency, pain is quantifiable (missed loads = lost revenue),",
         "accessible on LinkedIn, 5–20 truck operations can write a check without procurement,",
         "familiar with paying for operational tools.",
         "",
         "Second priority: Agriculture. Tight referral networks — one good client can multiply.",
         "Time the approach to their calendar (post-harvest or pre-season planning).",
         "",
         "Third priority: Field operations. Higher effort to reach, save for after initial traction."],
        style="info"
    )

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 15 — FOUNDER-LED SALES
    # ─────────────────────────────────────────────────────────────────────────
    add_section_heading(doc, "15", "FOUNDER-LED SALES CONVERSATION STRUCTURE")

    add_body(doc, "Wrong mindset: 'I need to convince them to buy.'", bold=True, colour=WARN_AMBER)
    add_body(doc, "Right mindset: 'I need to understand their operation well enough to know if I can actually help.'",
             bold=True, colour=ACCENT)

    add_h2(doc, "20-Minute Discovery Call Script")

    call_structure = [
        ("Minutes 0–2: Set the agenda",
         ["'Thanks for making time. I want to spend about 15 minutes understanding your operation —",
          "specifically how information flows and where decisions get slowed down. Then I will tell you",
          "what the diagnostic engagement looks like and you can tell me if it makes sense. Sound good?'"]),
        ("Minutes 2–8: Operational reality questions",
         ["'Tell me about your operation — size, how it is structured, what the busiest period looks like.'",
          "'When information needs to move from one part of your team to another, what happens? Walk me through a real example.'",
          "'What is the decision you most often wish you had better information for when you are making it?'",
          "'When something goes wrong operationally — a missed job, a scheduling clash — where does it usually break down?'"]),
        ("Minutes 8–12: Probe the cost",
         ["'When that breakdown happens, what does it actually cost — in time, in relationships, in revenue?'",
          "'How often does this happen?'",
          "'Have you tried to fix it before? What happened?'"]),
        ("Minutes 12–16: Offer description",
         ["'Based on what you have described, here is what the diagnostic looks like: [specific to what they said].",
          "The output is [specific deliverable]. It takes 2–3 weeks. Here is what it costs.'"]),
        ("Minutes 16–20: Temperature check",
         ["'Does that feel relevant to where you are right now?' — then stop talking.",
          "If yes: 'Can we set up the next step now?'",
          "If maybe: 'What would need to be true for it to be relevant?'",
          "If no: 'What would you actually need help with?' — they may tell you what to build."]),
    ]

    for step, bullets in call_structure:
        add_h3(doc, step)
        for b in bullets:
            add_bullet(doc, b, level=1)

    add_h2(doc, "High-Value Questions to Ask During Calls")
    hv_questions = [
        "Walk me through what happens between [X event] and [Y decision].",
        "Who does not have visibility into that, and what does that cost you?",
        "What is the last time this caused a real problem? What did it cost?",
        "If I could hand you one thing tomorrow morning that would change how you run operations — what would it be?",
        "What have you already tried to solve this?",
    ]
    for q in hv_questions:
        add_bullet(doc, q)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 16 — METRICS
    # ─────────────────────────────────────────────────────────────────────────
    add_section_heading(doc, "16", "WEEKLY METRICS TO TRACK")

    add_data_table(doc,
        ["Metric", "Target (Weeks 1–4)", "What It Tells You"],
        [
            ["Outreach sent",                  "15–25/week",      "Activity level"],
            ["Response rate",                  ">10%",            "Message quality"],
            ["Calls booked per 10 outreaches", ">1",              "Offer / CTA quality"],
            ["Calls completed",                "2–5/week",        "Conversion from interest"],
            ["Calls with pain confirmed (%)",  "Track all",       "ICP accuracy"],
            ["Proposals / offers made",        "1+ by Week 3",    "Revenue pipeline"],
            ["Paid engagements closed",        "1 by Week 4–6",   "Actual validation"],
            ["Pain patterns logged",           "Every call",      "Product research data"],
        ],
        col_widths=[2.8, 1.6, 2.3]
    )

    add_h2(doc, "What NOT to Track Right Now")
    for line in [
        "Website traffic",
        "LinkedIn impressions or follower growth",
        "Email open rates",
        "Time spent on product or positioning refinement",
    ]:
        add_bullet(doc, line, colour=WARN_AMBER)

    add_body(doc,
        "The only metric that matters in the next 30 days: did you have a real sales conversation?"
    )

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 17 — URGENCY
    # ─────────────────────────────────────────────────────────────────────────
    add_section_heading(doc, "17", "CREATING URGENCY WITHOUT MANIPULATION")

    add_h2(doc, "What Does NOT Work")
    for line in [
        "'Only accepting 3 clients this quarter' — you have zero clients",
        "'Prices going up next month' — arbitrary and transparent",
        "'Limited spots available' — meaningless before you are oversubscribed",
    ]:
        add_bullet(doc, line, colour=WARN_AMBER)

    add_h2(doc, "What Creates Legitimate Urgency")

    urgency_methods = [
        ("Connect to Their Operational Calendar",
         "You mentioned harvest season starts in 6 weeks. The diagnostic takes 2–3 weeks. "
         "If you want the output before peak season, we would need to start in the next 10 days. "
         "This is real urgency — tied to their reality, not yours."),
        ("Make the Cost of Delay Concrete",
         "Every month you are running on fragmented information, you are making decisions with "
         "incomplete data. That does not resolve on its own — it compounds."),
        ("Reference Their Own Incident",
         "You described the load coordination failure last month. That pattern keeps happening "
         "until the information flow is fixed. The diagnostic is how you start that."),
        ("The Pipeline Reality (only when true)",
         "I am currently working with [X] operation in your space. If you want to start this "
         "quarter, it makes sense to get the intake scheduled this week. Only say this when actually true."),
    ]

    for title, desc in urgency_methods:
        add_h3(doc, title)
        add_body(doc, desc, colour=MID_GREY)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 18 — MANUAL FIRST
    # ─────────────────────────────────────────────────────────────────────────
    add_section_heading(doc, "18", "WHAT TO BUILD MANUALLY FIRST")

    add_body(doc,
        "Everything is manual until a paying client needs it to be otherwise. "
        "The automation trap: every hour building systems is an hour not spent on "
        "conversations. You do not need automation. You need conversations."
    )

    add_data_table(doc,
        ["Thing", "Manual Version Now", "Build / Automate When"],
        [
            ["Intake form",          "Google Form (5 questions)",          "Never — works at scale"],
            ["Operations mapping",   "Google Slides or Figma manually",    "After 10+ diagnostics"],
            ["Diagnostic report",    "Google Docs template you fill in",   "After 5+ diagnostics"],
            ["Scheduling",           "Calendly free / Google Meet link",   "Already good enough"],
            ["CRM",                  "Google Sheet",                       "After 50+ contacts"],
            ["Proposal",             "PDF from a Google Docs template",    "Already good enough"],
            ["Invoicing",            "Stripe payment link or Wave",        "From first paid engagement"],
            ["Case studies",         "Write after first client",           "After first 2 clients"],
            ["Dashboard / platform", "Do not build",                       "After 5 paid diagnostics with consistent pattern"],
        ],
        col_widths=[1.8, 2.2, 2.7]
    )

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 19 — SCORING MODELS
    # ─────────────────────────────────────────────────────────────────────────
    add_section_heading(doc, "19", "HOW TO VALIDATE OPERATIONAL SCORING MODELS")

    add_callout(doc, "THE RISK",
        ["A scoring model built on 0–3 diagnostics reflects your assumptions, not operational reality.",
         "If you pitch it as methodology before it is validated, you will look naive to a sophisticated buyer",
         "or get locked into a model that does not hold."],
        style="warn"
    )

    add_h2(doc, "The Right Sequence")
    add_flow(doc, [
        "Do 3–5 paid diagnostics manually",
        "After each: what factors mattered most? What was highest-leverage?",
        "After 5 diagnostics: look for patterns — do 3–5 variables appear across all of them?",
        "Only then: codify into a scoring or assessment model",
        "Test the model on the next 3 diagnostics — does it hold?",
        "After 10 diagnostics: you have a methodology worth naming",
    ])

    add_h2(doc, "What to Say When Asked 'How Do You Measure Coordination Health?'")
    add_callout(doc, "RESPONSE",
        ["'We do not use a generic scorecard — the diagnostic produces a custom map of your operation.",
         "What we have found is that the critical variables are different enough across industries",
         "and company sizes that a single model does not hold. That is why we start with diagnosis",
         "rather than benchmarks.'",
         "",
         "This is honest AND differentiating. It signals depth, not a vendor selling",
         "a pre-packaged assessment."],
        style="info"
    )

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 20 — POSITIONING
    # ─────────────────────────────────────────────────────────────────────────
    add_section_heading(doc, "20", "POSITIONING vs. AI DASHBOARDS")

    add_body(doc,
        "Every prospect you will talk to has been pitched AI-powered dashboards, workflow tools, "
        "and 'operational intelligence' platforms. Your first job in every conversation is to break that pattern."
    )

    add_callout(doc, "CORE DISTINCTION",
        ["Software products make a bet on what your problem is.",
         "Norvane diagnoses what your problem actually is before recommending anything."],
        style="info"
    )

    add_h2(doc, "Positioning Statement Options")
    positions = [
        ("A — Diagnosis Before Prescription",
         "Most operations problems look like software problems from the outside. "
         "We do the diagnostic work to figure out what you actually need — which is often not software."),
        ("B — The Cost of the Wrong Fix",
         "The most expensive operations mistake is buying the right tool for the wrong problem. "
         "Norvane diagnoses the problem first."),
        ("C — Specificity Positioning",
         "We work specifically with ag, logistics, and field operations — not because we have a "
         "platform built for them, but because operational fragmentation shows up differently in "
         "each context and we have mapped those patterns."),
        ("D — Against AI Hype (when prospect has been burned by AI pitches)",
         "We do not sell AI-powered anything. We sit with your operations team, map where "
         "information breaks down, and tell you what to fix first. The output is a clear picture, not a platform."),
    ]

    for label, text in positions:
        add_h3(doc, label)
        add_body(doc, text, colour=MID_GREY)

    add_h2(doc, "The One-Sentence Test")
    add_body(doc,
        "Can your prospect describe what you do to a colleague accurately? "
        "'They do operations consulting' = too vague. "
        "'They diagnose where your operation loses information and what that is costing you' = "
        "specific enough to generate referrals."
    )

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 21 — PMF SIGNALS
    # ─────────────────────────────────────────────────────────────────────────
    add_section_heading(doc, "21", "SIGNALS OF REAL PRODUCT-MARKET FIT")

    add_h2(doc, "You Are On Track When")
    on_track = [
        "A prospect describes your service to their partner before you pitch them again",
        "Someone says 'we have been trying to figure out how to do exactly this'",
        "A client returns for a second engagement without you asking",
        "You hear the same operational problem described with the same language across 3 different companies",
        "A client credits the diagnostic with a specific decision they made differently",
        "Referrals come in without you asking",
    ]
    for line in on_track:
        add_bullet(doc, line, colour=ACCENT)

    add_h2(doc, "Reconsider Positioning When")
    reconsider = [
        "Every call requires a long explanation of what you do before they engage",
        "People understand the problem but not the engagement format",
        "Price objections happen before the value has been established",
        "You are attracting wrong-fit companies (wrong size, wrong industry, wrong pain)",
        "Nobody asks 'how much' or 'when can we start' after hearing the offer",
        "Calls are high-energy but never convert to agreed next steps",
    ]
    for line in reconsider:
        add_bullet(doc, line, colour=WARN_AMBER)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 22 — COMMON MISTAKES
    # ─────────────────────────────────────────────────────────────────────────
    add_section_heading(doc, "22", "COMMON FOUNDER MISTAKES AT THIS STAGE")

    mistakes = [
        "Polishing the website instead of doing outreach. The site is done. Every hour on it is productive procrastination.",
        "Making the diagnostic free 'to get proof.' Do not. Free work devalues the offer and attracts tire-kickers. Charge from engagement 1.",
        "Pitching too early on calls. The pitch comes after you understand the pain. If you are explaining Norvane before asking a single question, you are selling, not diagnosing.",
        "Building the methodology before doing the work. Say 'our diagnostic process,' not 'our proprietary framework.' One sounds grounded, the other sounds like consulting theater.",
        "Targeting enterprise for speed. Their procurement cycles are 3–6 months. You need 30 days. Stay in the 10–200 employee segment.",
        "Writing a full proposal before checking price sensitivity. Ask 'what is your budget for something like this?' on the call. Never write a 10-page proposal into the void.",
        "Being scared by 'we already have something like this.' This is not a rejection — it is an invitation to probe. Ask: what are they using, how well does it work, what does it miss?",
        "Needing to feel 'ready' before outreach. You are ready now. Start ugly.",
    ]

    for i, mistake in enumerate(mistakes, 1):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(3)
        para.paragraph_format.space_after  = Pt(3)
        para.paragraph_format.left_indent  = Inches(0.2)

        num_r = para.add_run(f"{i}.  ")
        num_r.font.name  = "Calibri"
        num_r.font.size  = Pt(10.5)
        num_r.font.bold  = True
        num_r.font.color.rgb = WARN_AMBER

        txt_r = para.add_run(mistake)
        txt_r.font.name  = "Calibri"
        txt_r.font.size  = Pt(10.5)
        txt_r.font.color.rgb = CHARCOAL

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 23 — WHAT NOT TO DO
    # ─────────────────────────────────────────────────────────────────────────
    add_section_heading(doc, "23", "WHAT NOT TO SPEND TIME ON RIGHT NOW")

    add_callout(doc, "THE 60-DAY FILTER RULE",
        ["Does this directly produce a conversation, a call, or a paid engagement within 60 days?",
         "Yes → do it.  No → do not."],
        style="warn"
    )

    add_h2(doc, "Hard Stops — Do Not Touch for 60 Days")
    stops = [
        "Building any software platform or dashboard",
        "Designing an operational scoring UI",
        "Writing detailed case studies before you have real clients",
        "A/B testing the website (not enough traffic to measure)",
        "SEO optimization",
        "Podcast, YouTube, or newsletter",
        "Hiring or finding partners",
        "Applying to accelerators or programs",
        "Writing detailed methodology whitepapers",
        "Brand guidelines documentation",
        "Fundraising conversations",
        "Enterprise outreach (Sysco, large ag corporations, freight networks) — wrong sales cycle",
        "Automating outreach before you know which manual message works",
    ]
    for s in stops:
        add_bullet(doc, s, colour=WARN_AMBER)

    # ─────────────────────────────────────────────────────────────────────────
    # APPENDIX — WEEK 1 CHECKLIST
    # ─────────────────────────────────────────────────────────────────────────
    page_break(doc)
    add_section_heading(doc, "APP", "WEEK 1 QUICK-START CHECKLIST")

    checklist_days = [
        ("Day 1", [
            "Build LinkedIn target list (50 people, logistics-first)",
            "Write 3 outreach message variants",
            "Send first 15 messages — 5 per variant",
        ]),
        ("Day 2", [
            "Publish first LinkedIn post (operational scenario + question)",
            "Follow up on Day 1 non-replies",
            "Set up Calendly free tier",
        ]),
        ("Day 3", [
            "Send second outreach wave (15 messages)",
            "Write call script — read aloud, time it",
            "Create 5-question intake form in Google Forms",
        ]),
        ("Days 4–6", [
            "Send outreach daily, follow up systematically",
            "Take every call that books — do not reschedule",
            "Log pain patterns after each call",
            "Publish second LinkedIn post",
        ]),
        ("End of Week 1", [
            "Review which message variant got responses",
            "Identify which industry and job title replied most",
            "Decide if ICP targeting adjustment is needed before Week 2",
        ]),
    ]

    for day, items in checklist_days:
        add_h2(doc, day)
        for item in items:
            add_bullet(doc, f"☐  {item}")

    # ─────────────────────────────────────────────────────────────────────────
    # CLOSING
    # ─────────────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph_border_bottom(closing, ACCENT)
    cr = closing.add_run("The only goal for the next 30 days: one paying client or enough clear rejections that you understand exactly what to change.")
    cr.font.name  = "Calibri"
    cr.font.size  = Pt(11)
    cr.font.bold  = True
    cr.font.italic = True
    cr.font.color.rgb = CHARCOAL

    doc.add_paragraph()
    closing2 = doc.add_paragraph()
    closing2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr2 = closing2.add_run("Everything else is a distraction.")
    cr2.font.name  = "Calibri"
    cr2.font.size  = Pt(12)
    cr2.font.bold  = True
    cr2.font.color.rgb = ACCENT

    # ── Save ──────────────────────────────────────────────────────────────────
    out_path = r"c:\Users\georg\OneDrive\Desktop\Projects\Norvane\Norvane_Operator_Playbook.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build()
